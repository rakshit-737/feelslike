"""Refresh docs/FeelsLike-Technical-Report.docx from the measured result files.

The long-form technical report is a hand-written Word document, not a generated
one, so this script does the same thing scripts/update_docs.py does for the deck:
it rewrites only the sentences, list items and table cells that carry MEASURED
numbers, leaving every other paragraph — and all of the document's styling —
untouched.

Editing technique: a paragraph's text is written into run 0 and the remaining
runs are blanked. Word splits these paragraphs into one run per word with
uniform formatting, so rebuilding a paragraph this way preserves its font,
weight and list level while making substring edits safe across run boundaries.

Usage:
  python -m scripts.update_technical_report --check    # report, write nothing
  python -m scripts.update_technical_report            # rewrite the .docx
  python -m scripts.update_technical_report --pdf      # also re-export the PDF
"""
from __future__ import annotations

import argparse
import json
import math
import re
import shutil
from pathlib import Path

ROOT = Path(__file__).resolve().parent.parent
DOCS = ROOT / "docs"
DOC = DOCS / "FeelsLike-Technical-Report.docx"


# ---------------------------------------------------------------------------
# measured facts
# ---------------------------------------------------------------------------

def wilson(hits: int, n: int, z: float = 1.96) -> tuple[int, int]:
    """95% Wilson score interval as whole percents.

    INPUT: successes, sample size. OUTPUT: (low_pct, high_pct).
    Wilson rather than the normal approximation because n is 20 and p is near
    the boundary, where the normal interval is visibly wrong.
    """
    if n == 0:
        return (0, 0)
    p = hits / n
    denom = 1 + z * z / n
    centre = (p + z * z / (2 * n)) / denom
    margin = z / denom * math.sqrt(p * (1 - p) / n + z * z / (4 * n * n))
    return (round(100 * max(0.0, centre - margin)), round(100 * min(1.0, centre + margin)))


def load_facts() -> dict:
    """Every number this document quotes, read from evals/*.json.

    INPUT: none. OUTPUT: dict of facts. SIDE EFFECTS: none.
    ERROR STATES: raises FileNotFoundError when results_nlp.json or
      results_blind.json is missing — both are load-bearing for this section, so
      failing loudly beats writing a half-updated document.
    """
    nlp = json.loads((ROOT / "evals/results_nlp.json").read_text())["splits"]
    blind = json.loads((ROOT / "evals/results_blind.json").read_text())

    def merge_heldout(splits: dict) -> dict:
        parts = [v for k, v in splits.items() if k.startswith("heldout")]
        keys = ("n", "det", "zone", "issue", "triple")
        return {k: sum(p.get(k, 0) for p in parts) for k in keys}

    facts = {
        "dev": nlp["dev"],
        "heldout": merge_heldout(nlp),
        "blind": blind,
    }
    facts["bench_n"] = facts["dev"]["n"] + facts["heldout"]["n"]

    llm_path = ROOT / "evals/results_nlp_llm.json"
    if llm_path.exists():
        lsplits = json.loads(llm_path.read_text())["splits"]
        facts["llm_dev"] = lsplits["dev"]
        facts["llm_heldout"] = merge_heldout(lsplits)

    blind_llm = ROOT / "evals/results_blind_llm.json"
    facts["blind_llm"] = json.loads(blind_llm.read_text()) if blind_llm.exists() else None
    return facts


# ---------------------------------------------------------------------------
# docx surgery
# ---------------------------------------------------------------------------

def set_para(par, new_text: str) -> bool:
    """Replace a paragraph's text, keeping run 0's formatting. True when changed."""
    if not par.runs or par.text.strip() == new_text.strip():
        return False
    par.runs[0].text = new_text
    for r in par.runs[1:]:
        r.text = ""
    return True


def set_cell(cell, new_text: str) -> bool:
    """Replace a table cell's text through its first paragraph."""
    par = cell.paragraphs[0]
    if not par.runs:
        return False
    return set_para(par, new_text)


def find_para(doc, predicate):
    """First paragraph whose text satisfies predicate(text), else None."""
    for p in doc.paragraphs:
        if p.text.strip() and predicate(p.text.strip()):
            return p
    return None


def edit(doc, matcher, build, label: str, log: list) -> bool:
    """Rewrite one paragraph located by matcher. build(old_text) -> new_text."""
    par = find_para(doc, matcher)
    if par is None:
        log.append(f"  MISS  {label}: no paragraph matched (document reworded?)")
        return False
    old = par.text.strip()
    new = build(old)
    if new is None:
        log.append(f"  same  {label}")
        return False
    if set_para(par, new):
        log.append(f"  edit  {label}: {old[:44]!r} -> {new[:44]!r}")
        return True
    log.append(f"  same  {label}: already current")
    return False


def pct_cell(hits: int, n: int) -> str:
    return f"{hits}/{n} ({100 * hits / n:.0f}%)" if n else "—"


def update_parser_table(doc, facts: dict, log: list) -> bool:
    """Rewrite Table 4 (parser accuracy) from the split results.

    The table is located by its header row rather than by index, so inserting a
    table earlier in the document cannot silently redirect the edit. Header cells
    carry their own sample size because the held-out block grew and the rules and
    LLM columns are no longer measured on the same number of cases.
    """
    dev, ho = facts["dev"], facts["heldout"]
    ldev, lho = facts.get("llm_dev"), facts.get("llm_heldout")
    for table in doc.tables:
        head = [c.text.strip().upper() for c in table.rows[0].cells]
        if not head or "METRIC" not in head[0] or len(table.columns) < 5:
            continue
        set_cell(table.rows[0].cells[1], f"RULES, DEV ({dev['n']})")
        set_cell(table.rows[0].cells[2], f"RULES, HELD-OUT ({ho['n']})")
        if ldev and lho:
            set_cell(table.rows[0].cells[3], f"LLM, DEV ({ldev['n']})")
            set_cell(table.rows[0].cells[4], f"LLM, HELD-OUT ({lho['n']})")
        rows = {"COMPLAINT DETECTION": "det", "ZONE EXTRACTION": "zone",
                "ISSUE EXTRACTION": "issue", "EXACT TRIPLE": "triple"}
        changed = 0
        for row in table.rows[1:]:
            key = rows.get(row.cells[0].text.strip().upper())
            if key is None:
                continue
            changed += set_cell(row.cells[1], pct_cell(dev[key], dev["n"]))
            changed += set_cell(row.cells[2], pct_cell(ho[key], ho["n"]))
            if ldev and lho:
                changed += set_cell(row.cells[3], pct_cell(ldev[key], ldev["n"]))
                changed += set_cell(row.cells[4], pct_cell(lho[key], lho["n"]))
        log.append(f"  edit  Table 4 (parser accuracy): {changed} cell(s) rewritten")
        return True
    log.append("  MISS  Table 4: no parser-accuracy table found")
    return False


def update_failure_list(doc, facts: dict, log: list) -> int:
    """Rewrite the failure-case bullets from the blind probe's own output.

    The previous bullets were held-out examples typed by hand; several of them
    have since been fixed, so the list described defects that no longer existed.
    Quoting the probe's failures directly makes that impossible.
    """
    fails = facts["blind"].get("failures") or []
    if not fails:
        return 0
    # Scoped to the failure-cases section: quoted bullets exist elsewhere in the
    # document, and an unscoped match would rewrite whichever one came first.
    paras = doc.paragraphs
    start = next((i for i, p in enumerate(paras)
                  if p.style.name.startswith("Heading")
                  and "failure cases" in p.text.strip().lower()), None)
    if start is None:
        log.append("  MISS  failure list: section heading not found")
        return 0
    end = next((i for i in range(start + 1, len(paras))
                if paras[i].style.name.startswith("Heading")), len(paras))
    bullets = [p for p in paras[start + 1:end]
               if p.style.name == "List Paragraph"
               and p.text.strip()[:1] in ("“", '"')]
    if not bullets:
        log.append("  MISS  failure list: no quoted bullets found")
        return 0
    written = 0
    for par, f in zip(bullets, fails):
        missed = ", ".join(f.get("missed", [])) or "the exact triple"
        line = (f"“{f['text']}” — {f.get('category', 'uncategorised')}. "
                f"The parser missed {missed}.")
        if set_para(par, line):
            written += 1
    # A shorter probe leaves stale bullets behind; drop them rather than keep them.
    for par in bullets[len(fails):]:
        par._element.getparent().remove(par._element)
        written += 1
    log.append(f"  edit  failure bullets: {written} rewritten from the blind probe")
    return written


def update_document(facts: dict, dry_run: bool) -> list:
    """Apply every measured-number edit the report needs. Returns the change log."""
    from docx import Document

    log: list = []
    doc = Document(str(DOC))

    dev, ho, b = facts["dev"], facts["heldout"], facts["blind"]
    bp, bt = b["pct"], b["tally"]
    lo, hi = wilson(bt["triple"], b["n"])
    bench_n = facts["bench_n"]
    llm_triple = (facts["blind_llm"] or {}).get("pct", {}).get("triple")

    edit(doc, lambda t: t.startswith("The benchmark has"),
         lambda _o: (
             f"The benchmark has {bench_n} labelled messages, scored on three fields "
             f"independently — is this a comfort complaint, which zone, which issue — "
             f"plus an “exact triple” score requiring all three to be right. The split "
             f"is the important part of the method. {dev['n']} cases are the development set the "
             f"rules parser was written against, so its perfect score there is a statement about "
             f"tuning, not about capability. A further {ho['n']} cases were written after the "
             f"rules were frozen: typos (“its friezing in the confrence room”), sarcasm, "
             f"Hinglish, multi-zone messages, retractions containing complaint keywords, and "
             f"negatives designed to defeat keyword matching (“the coffee machine is steaming "
             f"hot again”). Those cases have since been debugged against, which is exactly "
             f"what retires a split as a measure of generalisation. A separate {b['n']}-case "
             f"blind probe, never used to adjust anything, is the number we quote."),
         "benchmark method paragraph", log)

    gap = abs(bt["triple"] - round((llm_triple or 0) * b["n"] / 100)) if llm_triple else None
    comparison = ""
    if llm_triple is not None:
        comparison = (
            f" On the same probe the rules parser scores {bp['triple']}% exact triple against "
            f"the LLM path’s {llm_triple}% — a gap of {gap} message(s) on twenty, which "
            f"is a reason to keep measuring both paths rather than evidence that either one wins.")
    edit(doc, lambda t: t.startswith("Twenty cases is a small sample"),
         lambda _o: (
             f"Twenty cases is a small sample and we would rather say so than let the percentages "
             f"imply precision they do not have. The blind probe’s {bp['triple']}% exact "
             f"triple is {bt['triple']}/{b['n']}, which carries a 95% Wilson confidence interval "
             f"of roughly {lo}% to {hi}%.{comparison} What the numbers do support is the gap "
             f"between a tuned split and genuinely unseen phrasing, which is the whole reason the "
             f"probe exists."),
         "sample-size paragraph", log)

    update_parser_table(doc, facts, log)
    edit(doc, lambda t: t == "Failure cases from the held-out set",
         lambda _o: "Failure cases from the blind probe",
         "failure section heading", log)
    update_failure_list(doc, facts, log)

    edit(doc, lambda t: t.startswith("The third case is the one"),
         lambda _o: (
             "The false-positive cases are the ones that matter operationally, because they are "
             "the ones that reach the controller. The severity and decay mechanism bounds the "
             "damage (a single severity-1 or 2 offset that fades within an hour), but it does not "
             "prevent it."),
         "failure commentary", log)

    edit(doc, lambda t: t.startswith("A 20-case held-out benchmark")
         or t.startswith(f"The {b['n']}-case blind probe is small"),
         lambda _o: (
             f"The {b['n']}-case blind probe is small. Confidence intervals are wide: "
             f"{bp['triple']}% exact triple spans roughly {lo}–{hi}% at 95% (Wilson), so "
             f"differences of one or two messages are not differences at all."),
         "limitation: sample size", log)

    cats = [c for c, v in sorted(b.get("by_category", {}).items()) if v["triple"] < v["n"]]
    if cats:
        edit(doc, lambda t: t.startswith("Known parser failure modes remain"),
             lambda _o: ("Known parser failure modes remain. On the blind probe the categories "
                         "still failing are: " + ", ".join(cats) + "."),
             "limitation: failure modes", log)

    edit(doc, lambda t: t.startswith("The parts we are least sure of"),
         lambda o: re.sub(
             r"the parser’s [^,]+ accuracy is around 55–60%[^,]*|"
             r"the parser's [^,]+ accuracy is around 55–60%[^,]*|"
             r"the parser’s blind-probe accuracy is \d+% exact triple[^,]*|"
             r"the parser's blind-probe accuracy is \d+% exact triple[^,]*",
             f"the parser’s blind-probe accuracy is {bp['triple']}% exact triple on "
             f"deliberately difficult phrasing with a small sample behind it", o),
         "conclusion: parser accuracy", log)

    if not log:
        log.append("  document already matches every measured value")
    if not dry_run:
        backup = DOCS / (DOC.stem + ".backup.docx")
        shutil.copy2(DOC, backup)
        try:
            doc.save(str(DOC))
            log.append(f"  saved  {DOC.name} (backup: {backup.name})")
        except PermissionError:
            alt = DOCS / (DOC.stem + ".updated.docx")
            doc.save(str(alt))
            log.append(f"  LOCKED {DOC.name} is open in Word — wrote {alt.name} instead")
    return log


def export_pdf() -> str:
    """Re-export the PDF through Word. Returns a status line.

    Word is the only converter that reproduces this document's layout, so when
    it is unavailable the PDF is left alone and reported as stale rather than
    replaced with a differently-laid-out approximation.
    """
    try:
        import win32com.client  # type: ignore
    except ImportError:
        return ("  PDF NOT re-exported: pywin32 is not installed "
                "(pip install pywin32) — the .docx is current, the .pdf is not")
    out = DOCS / (DOC.stem + ".pdf")
    word = None
    try:
        word = win32com.client.Dispatch("Word.Application")
        word.Visible = False
        doc = word.Documents.Open(str(DOC), ReadOnly=1)
        doc.SaveAs(str(out), FileFormat=17)      # 17 = wdFormatPDF
        doc.Close(False)
        return f"  exported {out.name} ({out.stat().st_size // 1024} KB)"
    except Exception as exc:                      # Word absent or COM refused
        return f"  PDF NOT re-exported: {type(exc).__name__}: {exc}"
    finally:
        if word is not None:
            try:
                word.Quit()
            except Exception:
                pass


def main():
    ap = argparse.ArgumentParser()
    ap.add_argument("--check", action="store_true", help="report changes, write nothing")
    ap.add_argument("--pdf", action="store_true", help="re-export the PDF through Word")
    args = ap.parse_args()

    facts = load_facts()
    b = facts["blind"]
    print("MEASURED FACTS")
    print(f"  benchmark cases            {facts['bench_n']} "
          f"({facts['dev']['n']} dev + {facts['heldout']['n']} held-out)")
    print(f"  blind probe                {b['n']} cases, "
          f"{b['pct']['triple']}% exact triple, "
          f"95% CI {wilson(b['tally']['triple'], b['n'])}")

    print("\nDOCUMENT")
    for line in update_document(facts, args.check):
        print(line)
    if args.pdf and not args.check:
        print(export_pdf())


if __name__ == "__main__":
    main()
